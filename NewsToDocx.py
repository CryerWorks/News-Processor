from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

def format_markdown_bold(paragraph, text):
    """Detects Markdown bold (**bold**) and applies Word bold formatting."""
    pattern = r"\*\*(.*?)\*\*"  # Match text inside **double asterisks**
    matches = re.split(pattern, text)  # Split normal text and bold parts
    
    for i, part in enumerate(matches):
        run = paragraph.add_run(part)
        if i % 2 == 1:  # Every second part (inside ** **) should be bold
            run.bold = True
            
     # Apply consistent font styling to all text
        run.font.name = "Proxima Nova"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(53, 55, 68)  # Hex #353744

def add_auto_generated_toc(doc):
    toc_paragraph = doc.add_paragraph()
    run = toc_paragraph.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Proxima Nova"
    run.font.color.rgb = RGBColor(23, 84, 128)  # Hex #175480
    toc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Space below TOC
    
    # Insert Word's auto-generated TOC field
    toc = doc.add_paragraph()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = r'TOC \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    toc._element.append(fldChar1)
    toc._element.append(instrText)
    toc._element.append(fldChar2)
    toc._element.append(fldChar3)
    
    doc.add_paragraph("\n")  # Space after TOC

def convert_markdown_to_word(input_md, output_docx, logo_path):
    doc = Document()
    
    # Add front page with centered logo
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = logo_paragraph.add_run()
    run.add_picture(logo_path, width=Cm(14.67), height=Cm(7.32))
    
    title = doc.add_paragraph()
    title_run = title.add_run("Mundus Prime\nMonthly News Digest")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.name = "Proxima Nova"
    title_run.font.color.rgb = RGBColor(23, 84, 128)  # Hex #175480
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add auto-generated Table of Contents **immediately after title**
    add_auto_generated_toc(doc)
    
    doc.add_page_break()
    
    # Read markdown content
    with open(input_md, "r", encoding="utf-8") as file:
        lines = file.readlines()
    
    # Process content and structure it in Word
    for line in lines:
        line = line.strip()
        
        if line.startswith("# "):
            heading = doc.add_paragraph(line[2:], style="Heading 1")
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = heading.runs[0]
            run.font.name = "Proxima Nova"
            run.font.color.rgb = RGBColor(23, 84, 128)  # Hex #175480
        elif line.startswith("## "):
            heading = doc.add_paragraph(line[3:], style="Heading 2")
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = heading.runs[0]
            run.font.size = Pt(14)
            run.font.name = "Proxima Nova"
            run.font.color.rgb = RGBColor(23, 84, 128)  # Hex #175480
        elif line.startswith("### "):
            heading = doc.add_paragraph(line[4:], style="Heading 3")
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = heading.runs[0]
            run.font.name = "Proxima Nova"
            run.font.color.rgb = RGBColor(23, 84, 128)  # Hex #175480
        elif line:
            paragraph = doc.add_paragraph()
            format_markdown_bold(paragraph, line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = paragraph.runs[0]
            run.font.size = Pt(11)
            run.font.name = "Proxima Nova"
            run.font.color.rgb = RGBColor(53, 55, 68)  # Hex #353744
    
    # Save to Word document
    doc.save(output_docx)
    print(f"✅ Converted {input_md} to {output_docx}")

def main():
    """
    Main function that runs when the script is executed directly.
    This function is not called when the script is imported.
    """
    input_md = "Monthly_News_Digest.md"
    output_docx = "Monthly_News_Digest.docx"
    logo_path = "Mundus_Icon.png"
    
    convert_markdown_to_word(input_md, output_docx, logo_path)

if __name__ == "__main__":
    main()
